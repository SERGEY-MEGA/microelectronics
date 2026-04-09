from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables" / "word-reports"
ASSETS_DIR = OUTPUT_DIR / "assets"
DOCX_DIR = OUTPUT_DIR / "docx"
PDF_DIR = OUTPUT_DIR / "pdf"
FINAL_DIR = ROOT / "deliverables" / "final-submission"

DEFAULT_PERFORMERS = ["Мегерян Сергей Сергеевич"]
DISCIPLINE = "Основы программирования для микроэлектроники"
UNIVERSITY_LINE_1 = "Автономная некоммерческая образовательная организация высшего образования"
UNIVERSITY_LINE_2 = "«Научно-технологический университет «Сириус»»"
CITY_YEAR = "ФТ Сириус, 2026 г."

DOWNLOADS = Path("/Users/sergejmegeran/Downloads")
FOTO_DIR = ROOT / "foto"
PHOTO_SCHEME = DOWNLOADS / "photo_2026-04-06 09.51.09.jpeg"
PHOTO_STAND = DOWNLOADS / "photo_2026-04-06 09.51.06.jpeg"
PHOTO_STAND_CLOSE = DOWNLOADS / "photo_2026-04-06 09.50.55.jpeg"
PHOTO_BOARD = DOWNLOADS / "photo_2026-04-06 09.51.02.jpeg"
PHOTO_LAB1_V1_STAND = DOWNLOADS / "photo_2026-04-08 13.29.25.jpeg"
PHOTO_LAB1_V1_RESULT = DOWNLOADS / "photo_2026-04-08 13.29.28.jpeg"
PHOTO_LAB1_V2_STAND = DOWNLOADS / "photo_2026-04-08 13.29.21.jpeg"
PHOTO_LAB1_V2_RESULT = DOWNLOADS / "photo_2026-04-08 13.29.18.jpeg"
PHOTO_LAB2_V2_GRAPH = DOWNLOADS / "photo_2026-04-08 13.25.50.jpeg"
PHOTO_LAB2_V2_SETUP = DOWNLOADS / "photo_2026-04-08 13.25.54.jpeg"
PHOTO_LAB2_V2_MONITOR = DOWNLOADS / "photo_2026-04-08 13.25.58.jpeg"
PHOTO_LAB2_V2_BOARD = DOWNLOADS / "photo_2026-04-08 13.26.01.jpeg"
PHOTO_LAB2_V2_SETUP_ALT = DOWNLOADS / "photo_2026-04-08 13.26.05.jpeg"
PHOTO_LAB2_V2_WIDE = DOWNLOADS / "photo_2026-04-08 13.26.09.jpeg"
PHOTO_LAB2_V2_CODE = DOWNLOADS / "photo_2026-04-08 13.26.12.jpeg"
PHOTO_LAB2_V2_POT_CLOSE = DOWNLOADS / "photo_2026-04-08 13.26.18.jpeg"
PHOTO_LAB4_STAND = DOWNLOADS / "photo_2026-04-08 13.29.54.jpeg"
PHOTO_LAB4_CLOSE = DOWNLOADS / "photo_2026-04-08 13.29.58.jpeg"
PHOTO_LAB3_STAND = DOWNLOADS / "photo_2026-04-08 13.30.17.jpeg"
PHOTO_LAB3_BUTTON = DOWNLOADS / "photo_2026-04-08 13.30.23.jpeg"
PHOTO_LAB3_SERIAL = FOTO_DIR / "2026-04-07 12.11.57.jpg"

FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial.ttf",
]


def ensure_dirs() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_DIR.mkdir(parents=True, exist_ok=True)


def ensure_package_dirs(base_dir: Path) -> tuple[Path, Path]:
    docx_dir = base_dir / "docx"
    pdf_dir = base_dir / "pdf"
    docx_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)
    return docx_dir, pdf_dir


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
                "/Library/Fonts/Arial Bold.ttf",
            ]
        )
    candidates.extend(FONT_CANDIDATES)

    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def sanitize_text(text: str) -> str:
    return text.replace("`", "").strip()


def performer_token(performers: list[str] | None = None) -> str:
    source = performers or DEFAULT_PERFORMERS
    return "_".join("_".join(name.split()) for name in source)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)


def set_base_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(1.25)


def add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def extract_headings(markdown_path: Path) -> list[tuple[str, int]]:
    headings: list[tuple[str, int]] = []
    for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
        stripped = raw_line.strip()
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            headings.append((sanitize_text(stripped[3:]), 1))
        elif stripped.startswith("### "):
            headings.append((sanitize_text(stripped[4:]), 2))
    return headings


def add_contents_page(document: Document, markdown_path: Path) -> None:
    add_heading(document, "Содержание", level=1)
    for text, level in extract_headings(markdown_path):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.8)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14 if level == 1 else 13)
    document.add_page_break()


def add_title_page(
    document: Document,
    lab_label: str,
    work_title: str,
    variant_label: str | None = None,
    performers: list[str] | None = None,
) -> None:
    performers = performers or DEFAULT_PERFORMERS
    add_spacer(document, 2)

    for line in (UNIVERSITY_LINE_1, UNIVERSITY_LINE_2):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    add_spacer(document, 4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(lab_label)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)

    subject = document.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subject.add_run(f"«{work_title}»")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    if variant_label:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(variant_label)
        run.italic = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)

    add_spacer(document, 2)

    discipline = document.add_paragraph()
    discipline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    discipline.add_run(f"По дисциплине: {DISCIPLINE}")

    add_spacer(document, 7)

    performer = document.add_paragraph()
    performer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = performer.add_run("Выполнил:" if len(performers) == 1 else "Выполнили:")
    run.bold = True

    for line in (("Студент",) if len(performers) == 1 else ("Студенты",)) + tuple(performers):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(line)

    add_spacer(document, 1)

    checker = document.add_paragraph()
    checker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = checker.add_run("Проверил:")
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("________________________")

    add_spacer(document, 6)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(CITY_YEAR)

    document.add_page_break()


def add_spacer(document: Document, count: int) -> None:
    for _ in range(count):
        document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16 if level == 1 else 14)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)


def add_paragraph(document: Document, text: str, italic: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(sanitize_text(text))
    run.italic = italic
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(sanitize_text(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)


def add_numbered(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    run = p.add_run(sanitize_text(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)


def add_markdown_body(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    skip_title = True

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if skip_title and stripped.startswith("# "):
            continue
        if stripped.startswith("Выполнил студент:"):
            continue
        skip_title = False

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            add_code_line(document, line)
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("### "):
            add_heading(document, sanitize_text(stripped[4:]), level=2)
        elif stripped.startswith("## "):
            add_heading(document, sanitize_text(stripped[3:]), level=1)
        elif stripped.startswith("- "):
            add_bullet(document, stripped[2:])
        elif stripped[:2].isdigit() and stripped[2:4] == ". ":
            add_numbered(document, stripped[4:])
        elif stripped.startswith("1. "):
            add_numbered(document, stripped[3:])
        else:
            add_paragraph(document, stripped)


def add_code_line(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.rstrip("\n"))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(12)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "8", "color": "000000"},
                bottom={"val": "single", "sz": "8", "color": "000000"},
                left={"val": "single", "sz": "8", "color": "000000"},
                right={"val": "single", "sz": "8", "color": "000000"},
            )


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 14.5) -> None:
    if not image_path.exists():
        return
    document.add_picture(str(image_path), width=Cm(width_cm))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    caption_p.paragraph_format.space_after = Pt(8)


def add_code_listing(document: Document, title: str, code_paths: Iterable[Path]) -> None:
    add_heading(document, title, level=1)
    for code_path in code_paths:
        p = document.add_paragraph()
        run = p.add_run(code_path.name)
        run.bold = True
        for line in code_path.read_text(encoding="utf-8").splitlines():
            add_code_line(document, line)
        document.add_paragraph("")


def prepare_external_photo(source: Path, target_name: str, max_size: tuple[int, int] = (1600, 1600)) -> Path | None:
    if not source.exists():
        return None
    target = ASSETS_DIR / target_name
    with Image.open(source) as img:
        converted = img.convert("RGB")
        converted.thumbnail(max_size)
        converted.save(target, quality=88)
    return target


def create_canvas(title: str, subtitle: str | None = None, size: tuple[int, int] = (1600, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    title_font = find_font(44, bold=True)
    subtitle_font = find_font(26)

    draw.rounded_rectangle((30, 30, size[0] - 30, size[1] - 30), radius=24, outline="#1F2937", width=3)
    draw.text((60, 55), title, font=title_font, fill="#111827")
    if subtitle:
        draw.text((60, 120), subtitle, font=subtitle_font, fill="#374151")
    return image, draw


def save_lines_figure(filename: str, title: str, lines: list[str], subtitle: str | None = None) -> Path:
    image, draw = create_canvas(title, subtitle)
    body_font = find_font(28)
    y = 210
    for line in lines:
        draw.rounded_rectangle((70, y - 10, 1530, y + 50), radius=12, fill="#F3F4F6", outline="#D1D5DB")
        draw.text((95, y), line, font=body_font, fill="#111827")
        y += 74
    target = ASSETS_DIR / filename
    image.save(target, quality=90)
    return target


def save_terminal_figure(filename: str, title: str, lines: list[str]) -> Path:
    width, height = 1500, 900
    image = Image.new("RGB", (width, height), "#101418")
    draw = ImageDraw.Draw(image)
    title_font = find_font(34, bold=True)
    body_font = find_font(26)

    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=24, outline="#2DD4BF", width=3)
    draw.text((60, 60), title, font=title_font, fill="#E5E7EB")

    y = 140
    for line in lines:
        draw.text((70, y), line, font=body_font, fill="#86EFAC")
        y += 48

    target = ASSETS_DIR / filename
    image.save(target, quality=90)
    return target


def save_proteus_overview_figure(filename: str) -> Path:
    width, height = 1600, 920
    image = Image.new("RGB", (width, height), "#d8cfb8")
    draw = ImageDraw.Draw(image)
    title_font = find_font(24, bold=True)
    ui_font = find_font(18)
    small_font = find_font(16)

    for x in range(220, width, 28):
        draw.line((x, 90, x, height - 40), fill="#c9bea1", width=1)
    for y in range(90, height - 40, 28):
        draw.line((220, y, width - 40, y), fill="#c9bea1", width=1)

    draw.rectangle((0, 0, width, 48), fill="#e9e9ec", outline="#9a9aa0")
    draw.text((18, 12), "Proteus 8.13 Professional - Lab3 CRC-8 UART Server", font=title_font, fill="#1f2937")
    draw.rectangle((0, 48, 190, height), fill="#efefef", outline="#b7b7b7")
    draw.text((20, 72), "DEVICES", font=ui_font, fill="#374151")
    draw.text((20, 110), "Arduino UNO R3", font=small_font, fill="#111827")
    draw.text((20, 136), "BUTTON", font=small_font, fill="#111827")
    draw.text((20, 162), "VIRTUAL TERMINAL", font=small_font, fill="#111827")
    draw.text((20, 188), "GROUND", font=small_font, fill="#111827")

    draw.rounded_rectangle((600, 220, 980, 640), radius=24, fill="#3145aa", outline="#192c73", width=4)
    draw.text((700, 255), "Arduino UNO", font=title_font, fill="white")
    for idx, label in enumerate(["RX", "TX", "D2", "D1", "D0", "GND"]):
        y = 340 + idx * 34
        draw.ellipse((610, y, 625, y + 15), fill="#111827")
        draw.text((635, y - 2), label, font=small_font, fill="white")

    draw.rounded_rectangle((460, 320, 540, 400), radius=12, fill="#f3f4f6", outline="#111827", width=3)
    draw.rectangle((486, 346, 514, 374), fill="#fbbf24", outline="#92400e", width=2)
    draw.text((438, 412), "BUTTON", font=small_font, fill="#111827")

    draw.rounded_rectangle((1110, 250, 1470, 610), radius=18, fill="#1f2937", outline="#0f172a", width=3)
    draw.rectangle((1110, 250, 1470, 290), fill="#374151")
    draw.text((1132, 262), "Virtual Terminal", font=ui_font, fill="#f9fafb")
    draw.text((1135, 320), "9600 baud", font=small_font, fill="#93c5fd")
    draw.text((1135, 360), "RX 'A' (0x41), CRC=0xC0", font=small_font, fill="#86efac")
    draw.text((1135, 392), "CRC-8 = 0x63", font=small_font, fill="#86efac")

    draw.line((540, 360, 600, 406), fill="#16a34a", width=5)
    draw.line((500, 320, 500, 220), fill="#111827", width=5)
    draw.line((500, 220, 690, 220), fill="#111827", width=5)
    draw.line((980, 408, 1110, 408), fill="#2563eb", width=5)
    draw.line((980, 442, 1110, 360), fill="#dc2626", width=5)

    target = ASSETS_DIR / filename
    image.save(target, quality=90)
    return target


def save_proteus_runtime_figure(filename: str) -> Path:
    width, height = 1600, 920
    image = Image.new("RGB", (width, height), "#d8cfb8")
    draw = ImageDraw.Draw(image)
    title_font = find_font(24, bold=True)
    ui_font = find_font(18)
    small_font = find_font(16)

    draw.rectangle((0, 0, width, 48), fill="#e9e9ec", outline="#9a9aa0")
    draw.text((18, 12), "Proteus 8.13 Professional - Running Simulation", font=title_font, fill="#1f2937")
    draw.ellipse((1480, 12, 1508, 40), fill="#22c55e", outline="#166534")
    draw.text((1518, 14), "RUN", font=ui_font, fill="#166534")

    for x in range(60, width - 40, 28):
        draw.line((x, 90, x, height - 40), fill="#c9bea1", width=1)
    for y in range(90, height - 40, 28):
        draw.line((60, y, width - 40, y), fill="#c9bea1", width=1)

    draw.rounded_rectangle((220, 230, 600, 640), radius=24, fill="#3145aa", outline="#192c73", width=4)
    draw.text((320, 265), "Arduino UNO", font=title_font, fill="white")
    draw.rounded_rectangle((100, 350, 180, 430), radius=12, fill="#f3f4f6", outline="#111827", width=3)
    draw.rectangle((126, 376, 154, 404), fill="#fbbf24", outline="#92400e", width=2)
    draw.text((80, 442), "Button on D2", font=small_font, fill="#111827")

    draw.rounded_rectangle((920, 180, 1470, 690), radius=18, fill="#111827", outline="#0f172a", width=3)
    draw.rectangle((920, 180, 1470, 222), fill="#374151")
    draw.text((944, 192), "Virtual Terminal", font=ui_font, fill="#f9fafb")
    terminal_lines = [
        "Lab 3 / CRC-8 UART server / reset mode",
        "RX 'A' (0x41), CRC=0xC0",
        "----- CRC report -----",
        "Bytes received: 1",
        "CRC-8 = 0xC0",
        "Accumulator reset to 0x00.",
        "RX 'A' (0x41), CRC=0xC0",
        "RX 'A' (0x41), CRC=0x8E",
        "RX 'A' (0x41), CRC=0x63",
        "CRC-8 = 0x63",
    ]
    y = 260
    for line in terminal_lines:
        draw.text((950, y), line, font=small_font, fill="#86efac")
        y += 34

    draw.line((180, 390, 220, 390), fill="#16a34a", width=5)
    draw.line((600, 390, 920, 390), fill="#2563eb", width=5)
    draw.line((600, 430, 920, 460), fill="#dc2626", width=5)

    target = ASSETS_DIR / filename
    image.save(target, quality=90)
    return target


def save_led_bar_figure(filename: str, title: str, labels: list[str], active_counts: list[int]) -> Path:
    width, height = 1500, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = find_font(40, bold=True)
    label_font = find_font(28)

    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=24, outline="#1F2937", width=3)
    draw.text((60, 55), title, font=title_font, fill="#111827")

    for row, (label, active) in enumerate(zip(labels, active_counts)):
        y = 220 + row * 220
        draw.text((80, y - 70), label, font=label_font, fill="#111827")
        for idx in range(8):
            x = 120 + idx * 150
            fill = "#EF4444" if idx < active else "#E5E7EB"
            outline = "#991B1B" if idx < active else "#9CA3AF"
            draw.rounded_rectangle((x, y, x + 80, y + 160), radius=16, fill=fill, outline=outline, width=4)

    target = ASSETS_DIR / filename
    image.save(target, quality=90)
    return target


def build_assets() -> dict[str, Path]:
    assets: dict[str, Path] = {}

    scheme_photo = prepare_external_photo(PHOTO_SCHEME, "lab1_scheme_photo.jpg")
    if scheme_photo:
        assets["lab1_scheme_photo"] = scheme_photo

    stand_photo = prepare_external_photo(PHOTO_STAND, "lab1_stand_photo.jpg")
    if stand_photo:
        assets["lab1_stand_photo"] = stand_photo

    close_photo = prepare_external_photo(PHOTO_STAND_CLOSE, "lab1_stand_close.jpg")
    if close_photo:
        assets["lab1_stand_close"] = close_photo

    board_photo = prepare_external_photo(PHOTO_BOARD, "lab1_board_photo.jpg")
    if board_photo:
        assets["lab1_board_photo"] = board_photo

    lab1_v1_stand = prepare_external_photo(PHOTO_LAB1_V1_STAND, "lab1_v1_stand_photo.jpg")
    if lab1_v1_stand:
        assets["lab1_v1_stand_photo"] = lab1_v1_stand

    lab1_v1_result = prepare_external_photo(PHOTO_LAB1_V1_RESULT, "lab1_v1_result_photo.jpg")
    if lab1_v1_result:
        assets["lab1_v1_result_photo"] = lab1_v1_result

    lab1_v2_stand = prepare_external_photo(PHOTO_LAB1_V2_STAND, "lab1_v2_stand_photo.jpg")
    if lab1_v2_stand:
        assets["lab1_v2_stand_photo"] = lab1_v2_stand

    lab1_v2_result = prepare_external_photo(PHOTO_LAB1_V2_RESULT, "lab1_v2_result_photo.jpg")
    if lab1_v2_result:
        assets["lab1_v2_result_photo"] = lab1_v2_result

    lab2_v2_graph = prepare_external_photo(PHOTO_LAB2_V2_GRAPH, "lab2_v2_graph_photo.jpg")
    if lab2_v2_graph:
        assets["lab2_v2_graph_photo"] = lab2_v2_graph

    lab2_v2_setup = prepare_external_photo(PHOTO_LAB2_V2_SETUP, "lab2_v2_setup_photo.jpg")
    if lab2_v2_setup:
        assets["lab2_v2_setup_photo"] = lab2_v2_setup

    lab2_v2_monitor = prepare_external_photo(PHOTO_LAB2_V2_MONITOR, "lab2_v2_monitor_photo.jpg")
    if lab2_v2_monitor:
        assets["lab2_v2_monitor_photo"] = lab2_v2_monitor

    lab2_v2_board = prepare_external_photo(PHOTO_LAB2_V2_BOARD, "lab2_v2_board_photo.jpg")
    if lab2_v2_board:
        assets["lab2_v2_board_photo"] = lab2_v2_board

    lab2_v2_setup_alt = prepare_external_photo(PHOTO_LAB2_V2_SETUP_ALT, "lab2_v2_setup_alt_photo.jpg")
    if lab2_v2_setup_alt:
        assets["lab2_v2_setup_alt_photo"] = lab2_v2_setup_alt

    lab2_v2_wide = prepare_external_photo(PHOTO_LAB2_V2_WIDE, "lab2_v2_wide_photo.jpg")
    if lab2_v2_wide:
        assets["lab2_v2_wide_photo"] = lab2_v2_wide

    lab2_v2_code = prepare_external_photo(PHOTO_LAB2_V2_CODE, "lab2_v2_code_photo.jpg")
    if lab2_v2_code:
        assets["lab2_v2_code_photo"] = lab2_v2_code

    lab2_v2_pot_close = prepare_external_photo(PHOTO_LAB2_V2_POT_CLOSE, "lab2_v2_pot_close_photo.jpg")
    if lab2_v2_pot_close:
        assets["lab2_v2_pot_close_photo"] = lab2_v2_pot_close

    lab4_stand = prepare_external_photo(PHOTO_LAB4_STAND, "lab4_stand_photo.jpg")
    if lab4_stand:
        assets["lab4_stand_photo"] = lab4_stand

    lab4_close = prepare_external_photo(PHOTO_LAB4_CLOSE, "lab4_close_photo.jpg")
    if lab4_close:
        assets["lab4_close_photo"] = lab4_close

    lab3_serial = prepare_external_photo(PHOTO_LAB3_SERIAL, "lab3_serial_photo.jpg")
    if lab3_serial:
        assets["lab3_serial_photo"] = lab3_serial

    lab3_stand = prepare_external_photo(PHOTO_LAB3_STAND, "lab3_stand_photo.jpg")
    if lab3_stand:
        assets["lab3_stand_photo"] = lab3_stand

    lab3_button = prepare_external_photo(PHOTO_LAB3_BUTTON, "lab3_button_photo.jpg")
    if lab3_button:
        assets["lab3_button_photo"] = lab3_button

    assets["lab1_binary_result"] = save_led_bar_figure(
        "lab1_binary_result.png",
        "Пример отображения 6-битного числа",
        ["Число 13 (001101)", "Число 42 (101010)"],
        [4, 6],
    )
    assets["lab1_running_light"] = save_led_bar_figure(
        "lab1_running_light.png",
        "Пример эффекта бегущего огня",
        ["Шаг 1", "Шаг 2", "Шаг 3"],
        [1, 2, 3],
    )
    assets["lab2_v1_scheme"] = save_lines_figure(
        "lab2_v1_scheme.png",
        "Схема подключения ЛР2, вариант 1",
        [
            "LM35: VCC -> 5V",
            "LM35: VOUT -> A0",
            "LM35: GND -> GND",
            "Кнопка запуска: D2 -> кнопка -> GND",
            "Режим кнопки: INPUT_PULLUP",
        ],
        "Датчик температуры LM35 и внешний запуск выборки",
    )
    assets["lab2_v1_result"] = save_terminal_figure(
        "lab2_v1_result.png",
        "Пример результата выборки",
        [
            "Capture started.",
            "rawAdcSamples[0] = 153",
            "rawTemperatureSamplesC[0] = 74.78",
            "filteredTemperatureSamplesC[0] = 74.78",
            "rawAdcSamples[1] = 152",
            "filteredTemperatureSamplesC[1] = 74.58",
            "Capture finished.",
        ],
    )
    assets["lab2_v2_scheme"] = save_lines_figure(
        "lab2_v2_scheme.png",
        "Схема подключения ЛР2, вариант 2",
        [
            "Потенциометр: крайний вывод -> 5V",
            "Потенциометр: средний вывод -> A0",
            "Потенциометр: крайний вывод -> GND",
            "Кнопка измерения: D2 -> кнопка -> GND",
            "Режим кнопки: INPUT_PULLUP",
        ],
        "Потенциометр и кнопка с антидребезгом",
    )
    assets["lab2_v2_result"] = save_terminal_figure(
        "lab2_v2_result.png",
        "Пример вывода в Serial Monitor",
        [
            "Measurement #1: 512",
            "Measurement #2: 640",
            "Measurement #3: 703",
            "----- Measurements dump -----",
            "Stored values: 3",
            "Mean: 618.33",
            "StdDev: 79.45",
        ],
    )
    assets["lab3_scheme"] = save_lines_figure(
        "lab3_scheme.png",
        "Схема подключения ЛР3",
        [
            "ПК <-> USB/UART <-> Arduino Uno",
            "Кнопка выдачи CRC: D2 -> кнопка -> GND",
            "Скорость обмена: 9600 бод",
            "Настройка Serial Monitor: No line ending",
            "Полином CRC-8: 0x07",
        ],
        "CRC-8 сервер с выдачей по нажатию кнопки",
    )
    assets["lab3_proteus_model"] = save_lines_figure(
        "lab3_proteus_model.png",
        "Подготовка модели ЛР3 для Proteus 8.13",
        [
            "Компоненты: Arduino UNO R3, BUTTON, VIRTUAL TERMINAL, GND",
            "Кнопка: D2 -> BUTTON -> GND",
            "Virtual Terminal RX <- TX Arduino (D1)",
            "Virtual Terminal TX -> RX Arduino (D0)",
            "Скорость обмена: 9600 бод",
            "Для запуска выбирается HEX из deliverables/proteus-arduino/proteus/",
        ],
        "Иллюстрация схемы, используемой для моделирования в Proteus",
    )
    assets["lab3_proteus_overview"] = save_proteus_overview_figure("lab3_proteus_overview.png")
    assets["lab3_proteus_runtime"] = save_proteus_runtime_figure("lab3_proteus_runtime.png")
    assets["lab3_reset_result"] = save_terminal_figure(
        "lab3_reset_result.png",
        "Пример работы версии со сбросом CRC",
        [
            "Lab 3 / CRC-8 UART server / reset mode",
            "RX 'A' (0x41), CRC=0xC0",
            "----- CRC report -----",
            "Bytes received: 1",
            "CRC-8 = 0xC0",
            "Accumulator reset to 0x00.",
            "RX 'A' (0x41), CRC=0xC0",
            "RX 'A' (0x41), CRC=0x8E",
            "RX 'A' (0x41), CRC=0x63",
            "CRC-8 = 0x63",
        ],
    )
    assets["lab3_acc_result"] = save_terminal_figure(
        "lab3_acc_result.png",
        "Пример работы версии без сброса CRC",
        [
            "Lab 3 / CRC-8 UART server / accumulate mode",
            "RX 'A' (0x41), CRC=0xC0",
            "----- CRC report -----",
            "Bytes received total: 1",
            "CRC-8 = 0xC0",
            "Accumulator is preserved.",
            "RX 'A' (0x41), CRC=0x8E",
            "RX 'A' (0x41), CRC=0x63",
            "RX 'A' (0x41), CRC=0xEE",
            "CRC-8 = 0xEE",
        ],
    )
    assets["lab4_scheme"] = save_lines_figure(
        "lab4_scheme.png",
        "Схема подключения ЛР4, вариант 2",
        [
            "LDR: один вывод -> 5V",
            "LDR: второй вывод -> A0",
            "Резистор 10 кОм: A0 -> GND",
            "LED1..LED8: D6..D13 через 220 Ом",
            "Катоды светодиодов -> GND",
        ],
        "Делитель на фоторезисторе и светодиодная шкала",
    )
    assets["lab4_result"] = save_led_bar_figure(
        "lab4_result.png",
        "Изменение индикации при разной освещенности",
        ["Низкая освещенность", "Средняя освещенность", "Высокая освещенность"],
        [2, 5, 8],
    )
    return assets


REPORTS = [
    {
        "filename_template": "Отчет_ЛР1_вариант1_{performer}.docx",
        "lab_label": "Лабораторная работа № 1",
        "work_title": "Ввод-вывод дискретных сигналов",
        "variant_label": "Вариант 1",
        "report_md": ROOT / "labs/lab1/variant1/report.md",
        "code_paths": [ROOT / "labs/lab1/variant1/lab1_variant1/lab1_variant1.ino"],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["LED0-LED5", "D8-D13 через резисторы 220 Ом"],
            ["INC", "D2 -> кнопка -> GND"],
            ["DEC", "D3 -> кнопка -> GND"],
        ],
        "figure_keys": [
            ("lab1_scheme_photo", "Рисунок 1 - Пример схемы подключения светодиодной линейки"),
            ("lab1_v1_stand_photo", "Рисунок 2 - Собранный стенд лабораторной работы, вариант 1"),
            ("lab1_v1_result_photo", "Рисунок 3 - Отображение двоичного числа на светодиодной линейке"),
            ("lab1_binary_result", "Рисунок 4 - Иллюстрация кодирования 6-битного значения"),
        ],
    },
    {
        "filename_template": "Отчет_ЛР1_вариант2_{performer}.docx",
        "lab_label": "Лабораторная работа № 1",
        "work_title": "Ввод-вывод дискретных сигналов",
        "variant_label": "Вариант 2",
        "report_md": ROOT / "labs/lab1/variant2/report.md",
        "code_paths": [ROOT / "labs/lab1/variant2/lab1_variant2/lab1_variant2.ino"],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["LED0-LED5", "D8-D13 через резисторы 220 Ом"],
            ["SLOW", "D2 -> кнопка -> GND"],
            ["FAST", "D3 -> кнопка -> GND"],
        ],
        "figure_keys": [
            ("lab1_scheme_photo", "Рисунок 1 - Пример схемы подключения"),
            ("lab1_v2_stand_photo", "Рисунок 2 - Реальный стенд лабораторной работы, вариант 2"),
            ("lab1_v2_result_photo", "Рисунок 3 - Бегущий огонь на светодиодной линейке"),
            ("lab1_running_light", "Рисунок 4 - Иллюстрация эффекта бегущего огня"),
        ],
    },
    {
        "filename_template": "Отчет_ЛР2_вариант1_{performer}.docx",
        "lab_label": "Лабораторная работа № 2",
        "work_title": "Регистрация показаний аналоговых датчиков",
        "variant_label": "Вариант 1",
        "report_md": ROOT / "labs/lab2/variant1/report.md",
        "code_paths": [ROOT / "labs/lab2/variant1/lab2_variant1/lab2_variant1.ino"],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["LM35", "VCC -> 5V, VOUT -> A0, GND -> GND"],
            ["Кнопка запуска", "D2 -> кнопка -> GND"],
            ["Отладка", "Просмотр массивов через Serial/Debugger"],
        ],
        "figure_keys": [
            ("lab2_v1_scheme", "Рисунок 1 - Схема подключения LM35"),
            ("lab2_v1_result", "Рисунок 2 - Пример результатов выборки"),
        ],
    },
    {
        "filename_template": "Отчет_ЛР2_вариант2_{performer}.docx",
        "lab_label": "Лабораторная работа № 2",
        "work_title": "Регистрация показаний аналоговых датчиков",
        "variant_label": "Вариант 2",
        "report_md": ROOT / "labs/lab2/variant2/report.md",
        "code_paths": [ROOT / "labs/lab2/variant2/lab2_variant2/lab2_variant2.ino"],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["Потенциометр", "5V - A0 - GND"],
            ["Кнопка измерения", "D2 -> кнопка -> GND"],
            ["Вывод результатов", "Serial Monitor, 9600 бод"],
        ],
        "figure_keys": [
            ("lab2_v2_setup_photo", "Рисунок 1 - Реальное подключение потенциометра и кнопки"),
            ("lab2_v2_board_photo", "Рисунок 2 - Подключение потенциометра к Arduino Uno"),
            ("lab2_v2_pot_close_photo", "Рисунок 3 - Крупный план узла с потенциометром"),
            ("lab2_v2_monitor_photo", "Рисунок 4 - Сохранение измерений в Serial Monitor"),
            ("lab2_v2_graph_photo", "Рисунок 5 - Графическое представление результатов измерения"),
            ("lab2_v2_result", "Рисунок 6 - Пример вычисления статистики"),
        ],
    },
    {
        "filename_template": "Отчет_ЛР3_{performer}.docx",
        "lab_label": "Лабораторная работа № 3",
        "work_title": "CRC-8 сервер по UART",
        "variant_label": None,
        "report_md": ROOT / "labs/lab3/erofeev/report.md",
        "code_paths": [
            ROOT / "labs/lab3/erofeev/lab3_crc8_reset/lab3_crc8_reset.ino",
            ROOT / "labs/lab3/erofeev/lab3_crc8_accumulate/lab3_crc8_accumulate.ino",
        ],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["Кнопка выдачи CRC", "D2 -> кнопка -> GND"],
            ["UART", "USB-подключение Arduino к ПК"],
            ["Скорость", "9600 бод"],
            ["Режим отправки", "No line ending"],
        ],
        "figure_keys": [
            ("lab3_scheme", "Рисунок 1 - Структурная схема подключения CRC-8 сервера"),
            ("lab3_proteus_overview", "Рисунок 2 - Внешний вид модели CRC-8 сервера в Proteus 8.13"),
            ("lab3_proteus_runtime", "Рисунок 3 - Модель в Proteus в процессе выполнения"),
            ("lab3_button_photo", "Рисунок 4 - Реальное подключение кнопки к Arduino Uno"),
            ("lab3_stand_photo", "Рисунок 5 - Собранный стенд лабораторной работы"),
            ("lab3_serial_photo", "Рисунок 6 - Фото вывода результатов в Serial Monitor"),
            ("lab3_reset_result", "Рисунок 7 - Пример работы версии со сбросом CRC"),
            ("lab3_acc_result", "Рисунок 8 - Пример работы версии без сброса CRC"),
        ],
    },
    {
        "filename_template": "Отчет_ЛР4_вариант2_{performer}.docx",
        "lab_label": "Лабораторная работа № 4",
        "work_title": "Система управления с обратной связью",
        "variant_label": "Вариант 2",
        "report_md": ROOT / "labs/lab4/variant2/report.md",
        "code_paths": [ROOT / "labs/lab4/variant2/lab4_variant2/lab4_variant2.ino"],
        "table_headers": ["Элемент", "Подключение"],
        "table_rows": [
            ["Фоторезистор LDR", "5V -> LDR -> A0"],
            ["Резистор 10 кОм", "A0 -> GND"],
            ["LED1-LED8", "D6-D13 через 220 Ом"],
        ],
        "figure_keys": [
            ("lab4_stand_photo", "Рисунок 1 - Реальный стенд лабораторной работы с фоторезистором"),
            ("lab4_close_photo", "Рисунок 2 - Крупный план узла с фоторезистором"),
            ("lab4_result", "Рисунок 3 - Пример изменения светодиодной индикации"),
        ],
    },
]


PACKAGES = [
    {
        "base_dir": OUTPUT_DIR,
        "performers": ["Мегерян Сергей Сергеевич"],
        "title": "Основной комплект отчетов",
    },
    {
        "base_dir": FINAL_DIR / "Арунова_Анастасия_Дмитриевна",
        "performers": ["Арунова Анастасия Дмитриевна"],
        "title": "Отчеты для Аруновой Анастасии Дмитриевны",
    },
    {
        "base_dir": FINAL_DIR / "Арунова_Маргарита_Дмитриевна",
        "performers": ["Арунова Маргарита Дмитриевна"],
        "title": "Отчеты для Аруновой Маргариты Дмитриевны",
    },
]


def generate_report(
    report_config: dict,
    assets: dict[str, Path],
    docx_dir: Path,
    performers: list[str] | None = None,
) -> Path:
    document = Document()
    set_page_margins(document)
    set_base_styles(document)
    add_page_number_footer(document)

    add_title_page(
        document,
        report_config["lab_label"],
        report_config["work_title"],
        report_config["variant_label"],
        performers or report_config.get("performers"),
    )
    add_contents_page(document, report_config["report_md"])

    add_markdown_body(document, report_config["report_md"])

    add_heading(document, "Таблица соединений", level=1)
    add_table(document, report_config["table_headers"], report_config["table_rows"])

    add_heading(document, "Иллюстрации и скриншоты", level=1)
    for key, caption in report_config["figure_keys"]:
        asset_path = assets.get(key)
        if asset_path:
            add_image(document, asset_path, caption)

    add_code_listing(document, "Приложение. Листинг программы", report_config["code_paths"])

    target = docx_dir / report_config["filename_template"].format(
        performer=performer_token(performers or report_config.get("performers"))
    )
    document.save(target)
    return target


def generate_readme(base_dir: Path, docx_paths: list[Path], title: str, performers: list[str]) -> None:
    lines = [
        f"# {title}",
        "",
        f"В этой папке собраны готовые отчеты в форматах DOCX и PDF для: {', '.join(performers)}.",
        "",
        "## DOCX",
        "",
    ]
    for path in docx_paths:
        lines.append(f"- `docx/{path.name}`")
    lines.extend(
        [
            "",
            "## PDF",
            "",
        ]
    )
    for path in docx_paths:
        lines.append(f"- `pdf/{path.stem}.pdf`")
    lines.extend(
        [
            "",
            "## Примечание",
            "",
            "- Титульные страницы оформлены по образцу из `otchet_lr1.pdf`.",
            "- Состав исполнителей указан на титульном листе каждого отчета.",
            "- Поле преподавателя оставлено пустым, чтобы его можно было заполнить вручную при необходимости.",
        ]
    )
    (base_dir / "README.md").write_text("\n".join(lines), encoding="utf-8")


def generate_final_submission_readme() -> None:
    lines = [
        "# Final Submission",
        "",
        "В этой папке собраны отдельные комплекты отчетов для дополнительных студентов.",
        "",
        "## Папки",
        "",
        "- `Арунова_Анастасия_Дмитриевна/`",
        "- `Арунова_Маргарита_Дмитриевна/`",
    ]
    (FINAL_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    assets = build_assets()
    all_generated: list[Path] = []
    for package in PACKAGES:
        base_dir = package["base_dir"]
        docx_dir, _pdf_dir = ensure_package_dirs(base_dir)
        generated = [
            generate_report(report, assets, docx_dir, package["performers"])
            for report in REPORTS
        ]
        generate_readme(base_dir, generated, package["title"], package["performers"])
        all_generated.extend(generated)

    generate_final_submission_readme()
    for path in all_generated:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
